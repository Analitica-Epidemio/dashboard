"""
Generador de DOCX para boletines epidemiológicos.

Convierte contenido TipTap JSON a documentos Word (.docx) usando python-docx,
con gráficos renderizados como imágenes de alta resolución (300 DPI).
"""

import contextlib
import io
import json
import logging
import os
from typing import TYPE_CHECKING, Any, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.text.run import Run
    from sqlalchemy.ext.asyncio import AsyncSession

    from app.domains.boletines.models import BoletinInstance

logger = logging.getLogger(__name__)

# ── Font detection ───────────────────────────────────────────────────────
_GEIST_FONT_DIR = "/usr/share/fonts/truetype/geist"
_ROBOTO_FONT_DIR = "/usr/share/fonts/truetype/roboto"

_COLOR_DARK = RGBColor(0x11, 0x18, 0x27)  # #111827
_COLOR_GRAY = RGBColor(0x37, 0x41, 0x51)  # #374151
_COLOR_HEADER_BG = "F3F4F6"

# Heading sizes matching the PDF CSS
_HEADING_SIZES = {1: 22, 2: 18, 3: 15, 4: 13}


def _detect_font_family() -> str:
    """Detect best available font. Priority: Geist > Roboto > Calibri."""
    for family, font_dir in [("Geist", _GEIST_FONT_DIR), ("Roboto", _ROBOTO_FONT_DIR)]:
        if os.path.isdir(font_dir):
            logger.info("DOCX font: '%s'", family)
            return family
    logger.info("DOCX font: Calibri (fallback)")
    return "Calibri"


_FONT_FAMILY = _detect_font_family()


class BulletinDocxGenerator:
    """Genera documentos DOCX a partir de contenido TipTap JSON de boletines."""

    def __init__(
        self,
        db: "AsyncSession",
        instance: Optional["BoletinInstance"] = None,
    ):
        self.db = db
        self.instance = instance
        self.doc = Document()
        self._setup_document()
        self._configure_styles()

    # ── Setup ────────────────────────────────────────────────────────────

    def _setup_document(self) -> None:
        """Configure A4 page with 25.4mm margins."""
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)

    def _configure_styles(self) -> None:
        """Set default font and heading styles matching the PDF output."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = _FONT_FAMILY
        font.size = Pt(11)
        font.color.rgb = _COLOR_DARK
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

        for level in range(1, 5):
            heading_style = self.doc.styles[f"Heading {level}"]
            hf = heading_style.font
            hf.name = _FONT_FAMILY
            hf.size = Pt(_HEADING_SIZES[level])
            hf.color.rgb = _COLOR_DARK
            hf.bold = level == 4  # Only H4 is bold per CSS (font-weight: 600)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    # ── Main entry ───────────────────────────────────────────────────────

    async def generate_docx(self, content: str) -> bytes:
        """
        Generate DOCX bytes from instance content.

        Args:
            content: TipTap JSON string or HTML string.

        Returns:
            bytes of the generated .docx file.
        """
        content = content.strip()

        # Try to parse as TipTap JSON
        if content.startswith(("{", "[")):
            try:
                parsed = json.loads(content)
                if isinstance(parsed, dict) and parsed.get("type") == "doc":
                    await self._process_tiptap_doc(parsed)
                    return self._save_to_bytes()
            except json.JSONDecodeError:
                pass

        # Fallback: convert HTML to TipTap isn't possible,
        # render via HTML pipeline (convert to HTML first, then use tiptap)
        # For HTML content, convert through tiptap_to_html in reverse isn't feasible.
        # Instead, generate a simple paragraph with the raw content as a fallback.
        logger.warning("Content is not TipTap JSON, adding as raw text fallback")
        self.doc.add_paragraph(content)
        return self._save_to_bytes()

    def _save_to_bytes(self) -> bytes:
        """Save document to bytes buffer."""
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()

    # ── TipTap document processing ───────────────────────────────────────

    async def _process_tiptap_doc(self, doc: dict[str, Any]) -> None:
        """Walk the TipTap doc tree and build DOCX content."""
        nodes = doc.get("content", [])
        # Remove the default empty paragraph that python-docx adds
        if self.doc.paragraphs:
            p = self.doc.paragraphs[0]
            p_element = p._element
            p_element.getparent().remove(p_element)
        await self._process_nodes(nodes)

    async def _process_nodes(self, nodes: list[dict[str, Any]]) -> None:
        """Process a list of TipTap nodes."""
        for node in nodes:
            await self._process_node(node)

    async def _process_node(self, node: dict[str, Any]) -> None:
        """Dispatch a single TipTap node to the appropriate handler."""
        node_type = node.get("type", "")
        handler = getattr(self, f"_handle_{node_type}", None)
        if handler:
            await handler(node)
        elif node_type == "text":
            # Text nodes are handled inline by _add_inline_content
            pass
        else:
            # Unknown node: try to process children
            content = node.get("content", [])
            if content:
                await self._process_nodes(content)

    # ── Node handlers ────────────────────────────────────────────────────

    async def _handle_paragraph(self, node: dict[str, Any]) -> None:
        content = node.get("content", [])
        p = self.doc.add_paragraph()
        self._add_inline_content(p, content)

    async def _handle_heading(self, node: dict[str, Any]) -> None:
        level = node.get("attrs", {}).get("level", 1)
        content = node.get("content", [])
        p = self.doc.add_heading(level=level)
        self._add_inline_content(p, content)

    async def _handle_bulletList(self, node: dict[str, Any]) -> None:
        for item in node.get("content", []):
            await self._handle_listItem(item, style="List Bullet")

    async def _handle_orderedList(self, node: dict[str, Any]) -> None:
        for item in node.get("content", []):
            await self._handle_listItem(item, style="List Number")

    async def _handle_listItem(
        self, node: dict[str, Any], style: str = "List Bullet"
    ) -> None:
        content = node.get("content", [])
        for child in content:
            child_type = child.get("type", "")
            if child_type == "paragraph":
                p = self.doc.add_paragraph(style=style)
                self._add_inline_content(p, child.get("content", []))
            elif child_type in ("bulletList", "orderedList"):
                # Nested list — process recursively
                await self._process_node(child)
            else:
                p = self.doc.add_paragraph(style=style)
                self._add_inline_content(p, child.get("content", []))

    async def _handle_blockquote(self, node: dict[str, Any]) -> None:
        for child in node.get("content", []):
            if child.get("type") == "paragraph":
                p = self.doc.add_paragraph(style="Quote")
                self._add_inline_content(p, child.get("content", []))
            else:
                await self._process_node(child)

    async def _handle_codeBlock(self, node: dict[str, Any]) -> None:
        content = node.get("content", [])
        text = "".join(n.get("text", "") for n in content if n.get("type") == "text")
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        # Light gray background via shading
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F3F4F6")
        shading.set(qn("w:val"), "clear")
        p.paragraph_format.element.get_or_add_pPr().append(shading)

    async def _handle_horizontalRule(self, _node: dict[str, Any]) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border via XML
        pPr = p.paragraph_format.element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D1D5DB")
        pBdr.append(bottom)
        pPr.append(pBdr)

    async def _handle_image(self, node: dict[str, Any]) -> None:
        attrs = node.get("attrs", {})
        src = attrs.get("src", "")
        if src.startswith("data:"):
            # base64 data URI
            import base64

            try:
                # Extract base64 data after the comma
                b64_data = src.split(",", 1)[1]
                img_bytes = base64.b64decode(b64_data)
                self.doc.add_picture(io.BytesIO(img_bytes), width=Mm(160))
            except Exception:
                logger.warning("Failed to decode base64 image, skipping")
        else:
            # External URL – skip (can't embed easily)
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Imagen: {src}]")
            run.font.color.rgb = _COLOR_GRAY

    async def _handle_hardBreak(self, _node: dict[str, Any]) -> None:
        # Hard breaks are handled within inline content
        pass

    async def _handle_table(self, node: dict[str, Any]) -> None:
        rows_data = node.get("content", [])
        if not rows_data:
            return

        # Determine column count from the first row
        first_row = rows_data[0].get("content", [])
        col_count = sum(
            c.get("attrs", {}).get("colspan", 1) for c in first_row
        )
        if col_count == 0:
            col_count = len(first_row)

        table = self.doc.add_table(rows=0, cols=col_count)
        table.style = self.doc.styles["Table Grid"]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_node in rows_data:
            cells_data = row_node.get("content", [])
            row = table.add_row()

            col_idx = 0
            for cell_node in cells_data:
                if col_idx >= col_count:
                    break
                cell_type = cell_node.get("type", "")
                cell_attrs = cell_node.get("attrs", {})
                colspan = cell_attrs.get("colspan", 1)

                cell = row.cells[col_idx]

                # Merge for colspan
                if colspan > 1 and col_idx + colspan <= col_count:
                    cell = cell.merge(row.cells[col_idx + colspan - 1])

                # Write cell content
                cell_content = cell_node.get("content", [])
                # Clear default paragraph
                if cell.paragraphs:
                    cell.paragraphs[0].clear()
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()

                for child in cell_content:
                    if child.get("type") == "paragraph":
                        self._add_inline_content(p, child.get("content", []))
                    else:
                        # For non-paragraph content, just add text
                        text_parts = self._extract_text(child)
                        if text_parts:
                            p.add_run(text_parts)

                # Header cell styling
                if cell_type == "tableHeader":
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                    # Shade header cells
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), _COLOR_HEADER_BG)
                    shading.set(qn("w:val"), "clear")
                    tcPr.append(shading)

                col_idx += colspan

    async def _handle_dynamicChart(self, node: dict[str, Any]) -> None:
        """Render a dynamic chart as a 300 DPI image embedded in the document."""
        from app.domains.boletines.html_renderer import _resolve_chart_code
        from app.domains.charts.schemas import CodigoGrafico, FiltrosGrafico
        from app.domains.charts.services.renderer import chart_renderer
        from app.domains.charts.services.spec_generator import ChartSpecGenerator

        attrs = node.get("attrs", {})
        chart_code_raw: str = attrs.get("chartCode", "")
        title: str = attrs.get("title", "")

        if not chart_code_raw:
            logger.warning("dynamicChart without chartCode, skipping")
            return

        try:
            chart_code: CodigoGrafico = _resolve_chart_code(chart_code_raw)

            grupo_ids = self._parse_ids_attr(attrs.get("grupoIds", ""))
            evento_ids = self._parse_ids_attr(attrs.get("eventoIds", ""))
            fecha_desde: str | None = attrs.get("fechaDesde", "") or None
            fecha_hasta: str | None = attrs.get("fechaHasta", "") or None

            filtros = FiltrosGrafico(
                ids_grupo_eno=grupo_ids if grupo_ids else None,
                ids_tipo_eno=evento_ids if evento_ids else None,
                fecha_desde=fecha_desde,
                fecha_hasta=fecha_hasta,
            )

            generator = ChartSpecGenerator(self.db)
            spec = await generator.generar_spec(
                codigo_grafico=chart_code,
                filtros=filtros,
                configuracion={"height": 400},
            )

            img_bytes = chart_renderer.renderizar_a_bytes(spec, dpi=300)

            # Add optional title
            if title:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

            self.doc.add_picture(io.BytesIO(img_bytes), width=Mm(160))
            # Center the picture
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logger.info("Chart '%s' embedded in DOCX", chart_code)

        except Exception as e:
            logger.error("Error rendering chart '%s' for DOCX: %s", chart_code_raw, e)
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Error al generar gráfico: {chart_code_raw or 'desconocido'}]")
            run.font.color.rgb = _COLOR_GRAY

    async def _handle_pageBreak(self, _node: dict[str, Any]) -> None:
        self.doc.add_page_break()

    async def _handle_dynamicTable(self, node: dict[str, Any]) -> None:
        """Insert placeholder text for dynamic tables (same as PDF)."""
        attrs = node.get("attrs", {})
        title = attrs.get("title", "")
        query_type = attrs.get("queryType", "")

        table_labels = {
            "top_enos": "Eventos más frecuentes",
            "capacidad_hospitalaria": "Capacidad Hospitalaria",
            "casos_suh": "Casos de SUH",
        }
        table_name = table_labels.get(query_type, query_type)

        if title:
            p = self.doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(11)

        p = self.doc.add_paragraph()
        run = p.add_run(f"[Tabla dinámica: {table_name}]")
        run.font.color.rgb = _COLOR_GRAY
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    async def _handle_variableNode(self, node: dict[str, Any]) -> None:
        """Substitute variable with actual value from instance metadata."""
        attrs = node.get("attrs", {})
        var_key = attrs.get("variableKey", "")

        if not var_key:
            return

        variables = self._build_variable_values()
        value = variables.get(var_key)

        p = self.doc.add_paragraph()
        if value is not None:
            run = p.add_run(str(value))
            run.bold = True
        else:
            run = p.add_run(f"{{{{ {var_key} }}}}")
            run.font.color.rgb = _COLOR_GRAY

    # ── Inline content processing ────────────────────────────────────────

    def _add_inline_content(
        self, paragraph: "DocxParagraph", nodes: list[dict[str, Any]]
    ) -> None:
        """Process inline text nodes with marks and add to paragraph."""
        for node in nodes:
            node_type = node.get("type", "")

            if node_type == "text":
                text = node.get("text", "")
                marks = node.get("marks", [])
                run = paragraph.add_run(text)
                run.font.name = _FONT_FAMILY
                run.font.size = Pt(11)
                self._apply_marks(run, marks)

            elif node_type == "hardBreak":
                run = paragraph.add_run()
                run.add_break()

            elif node_type == "variableNode":
                # Inline variable substitution
                attrs = node.get("attrs", {})
                var_key = attrs.get("variableKey", "")
                variables = self._build_variable_values()
                value = variables.get(var_key)
                if value is not None:
                    run = paragraph.add_run(str(value))
                    run.bold = True
                    run.font.name = _FONT_FAMILY
                    run.font.size = Pt(11)
                else:
                    run = paragraph.add_run(f"{{{{ {var_key} }}}}")
                    run.font.name = _FONT_FAMILY
                    run.font.size = Pt(11)
                    run.font.color.rgb = _COLOR_GRAY

    def _apply_marks(self, run: "Run", marks: list[dict[str, Any]]) -> None:
        """Apply TipTap marks (bold, italic, etc.) to a python-docx Run."""
        for mark in marks:
            mark_type = mark.get("type", "")
            attrs = mark.get("attrs", {})

            if mark_type == "bold":
                run.bold = True
            elif mark_type == "italic":
                run.italic = True
            elif mark_type == "underline":
                run.underline = True
            elif mark_type == "strike":
                run.font.strike = True
            elif mark_type == "code":
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif mark_type == "link":
                # python-docx doesn't have easy hyperlink support on runs,
                # just underline + color to indicate it's a link
                run.underline = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # blue
            elif mark_type == "textStyle":
                color = attrs.get("color", "")
                if color and color.startswith("#") and len(color) == 7:
                    try:
                        r = int(color[1:3], 16)
                        g = int(color[3:5], 16)
                        b = int(color[5:7], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except ValueError:
                        pass
            elif mark_type == "highlight":
                color = attrs.get("color", "yellow")
                # Map common highlight colors to Word highlight
                from docx.enum.text import WD_COLOR_INDEX

                highlight_map = {
                    "yellow": WD_COLOR_INDEX.YELLOW,
                    "#FFFF00": WD_COLOR_INDEX.YELLOW,
                    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
                    "red": WD_COLOR_INDEX.RED,
                    "blue": WD_COLOR_INDEX.BLUE,
                }
                hl_val = highlight_map.get(color)
                if hl_val is not None:
                    run.font.highlight_color = hl_val

    # ── Helpers ──────────────────────────────────────────────────────────

    def _build_variable_values(self) -> dict[str, Any]:
        """Build variable values from instance metadata (same logic as html_renderer)."""
        if not self.instance:
            return {}

        inst = self.instance
        values: dict[str, Any] = {}

        if inst.anio_epidemiologico:
            values["anio_epidemiologico"] = str(inst.anio_epidemiologico)
        if inst.semana_epidemiologica:
            values["semana_epidemiologica_actual"] = str(inst.semana_epidemiologica)
        if inst.fecha_inicio:
            values["fecha_inicio_semana_epidemiologica"] = inst.fecha_inicio.strftime(
                "%d/%m/%Y"
            )
        if inst.fecha_fin:
            values["fecha_fin_semana_epidemiologica"] = inst.fecha_fin.strftime(
                "%d/%m/%Y"
            )
        if inst.num_semanas:
            values["num_semanas_analizadas"] = str(inst.num_semanas)
            if inst.semana_epidemiologica:
                se_inicio = inst.semana_epidemiologica - inst.num_semanas + 1
                values["semana_epidemiologica_inicio"] = str(max(1, se_inicio))

        params = inst.parameters or {}
        for key, val in params.items():
            if key not in values and val is not None:
                values[key] = str(val)

        return values

    @staticmethod
    def _parse_ids_attr(value: Any) -> list[int]:
        """Parse comma-separated IDs string to list of integers."""
        if not value:
            return []
        if isinstance(value, list):
            return [int(x) for x in value if x]
        result = []
        for x in str(value).split(","):
            x = x.strip()
            if x:
                with contextlib.suppress(ValueError):
                    result.append(int(x))
        return result

    @staticmethod
    def _extract_text(node: dict[str, Any]) -> str:
        """Recursively extract plain text from a node tree."""
        if node.get("type") == "text":
            return node.get("text", "")
        parts = [
            BulletinDocxGenerator._extract_text(child)
            for child in node.get("content", [])
        ]
        return "".join(parts)
